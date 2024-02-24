import math
from copy import deepcopy
from docx import Document
from docx.shared import RGBColor

from font import (
    set_table_font,
    color_filled_cells,
    oboji_tekst_celija_tabele,
    set_table_color,
)

doc = Document("tables_learning.docx")

list_tables = doc.tables

wallet_numbers = [
    "wallet_num=111-111",
    "wallet_num=222-222",
    "wallet_num=333-333",
    "wallet_num=444-444",
    "wallet_num=555-555",
    "wallet_num=777-777",
    "wallet_num=888-888",
    "wallet_num=999-999",
]

wallet_names = [
    "wallet_name1",
    "wallet_name2",
    "wallet_name3",
    "wallet_name4",
    "wallet_name5",
    "wallet_name6",
    "wallet_name7",
    "wallet_name8",
]

amounts = [
    "10.000",
    "20.000",
    "30.000",
    "40.000",
    "50.000",
    "60.000",
    "70.000",
    "10.000",
]

currencies = ["EUR", "DOLLAR", "RSD", "GBP", "EUR", "DOLLAR", "RSD", "EUR"]

index_wallet_number_cell = 0
index_wallet_name_cell = 0

original_table = doc.tables[0]

need_table = math.ceil(len(wallet_numbers) / 2)


def create_tables(need_table_for_create: int) -> list:
    num = 0
    while num < need_table_for_create:
        new_table_element = deepcopy(original_table._element)

        new_table = doc.add_table(
            rows=len(original_table.rows), cols=len(original_table.columns)
        )
        new_table._element.getparent().replace(new_table._element, new_table_element)
        num += 1

        line_breaks = 1
        for _ in range(line_breaks):
            doc.add_paragraph()

    tr = doc.tables[-1]._element
    tbl = tr.getparent()
    tbl.remove(tr)
    doc.save("output.docx")
    return doc.tables


def populate_wallet_name(tables: list, wallet_names: list):
    index_wallet_name_cell = 0
    step = 1
    end = len(wallet_names)
    start = 0
    for table in tables:
        for name in wallet_names[start:end:step]:
            wallet_name_cell = table.cell(0, index_wallet_name_cell)
            wallet_name_cell.text = name
            index_wallet_name_cell += 3
            if index_wallet_name_cell > 4:
                index_wallet_name_cell = 0
                break
        start += 2


def populate_wallet_number(tables: list, wallet_numbers: list):
    index_wallet_number_cell = 0
    step = 1
    end = len(wallet_names)
    start = 0
    for table in tables:
        for number in wallet_numbers[start:end:step]:
            wallet_number_cell = table.cell(1, index_wallet_number_cell)
            wallet_number_cell.text = number
            index_wallet_number_cell += 3
            if index_wallet_number_cell > 4:
                index_wallet_number_cell = 0
                break
        start += 2


def populate_amounts(tables: list, amounts: list):
    index_amount_cell = 0
    step = 1
    end = len(wallet_names)
    start = 0
    for table in tables:
        for amount in amounts[start:end:step]:
            amount_number_cell = table.cell(3, index_amount_cell)
            amount_number_cell.text = amount
            index_amount_cell += 3
            if index_amount_cell > 4:
                index_amount_cell = 0
                break
        start += 2


def populate_currencies(tables: list, currencies: list):
    index_currency_cell = 1
    step = 1
    end = len(currencies)
    start = 0
    for table in tables:
        for currency in currencies[start:end:step]:
            currency_number_cell = table.cell(3, index_currency_cell)
            currency_number_cell.text = currency
            index_currency_cell += 3
            if index_currency_cell > 5:
                index_currency_cell = 1
                break
        start += 2


def set_font_on_table(tables: list):
    for table in tables:
        set_table_font(table, "Arial", 9, RGBColor(0, 0, 0))


def set_color_on_table(tables: list):
    for table in tables:
        set_table_color(table)


if __name__ == "__main__":
    tables = create_tables(need_table)

    populate_wallet_name(tables, wallet_names)
    populate_wallet_number(tables, wallet_numbers)
    populate_amounts(tables, amounts)
    populate_currencies(tables, currencies)
    set_font_on_table(tables)
    set_color_on_table(tables)
    doc.save("output.docx")
